from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUT_FILE = Path("Herbal_Vision_项目设计文档.docx")


def set_run_font(run, name: str = "Microsoft YaHei", size: int | None = None, bold: bool | None = None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rpr.append(rfonts)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_run_font(r, size={1: 18, 2: 15, 3: 13}.get(level, 12), bold=True)
    return p


def add_para(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=11)


def build_doc() -> Document:
    doc = Document()

    # Global font defaults
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Herbal Vision 项目设计文档")
    set_run_font(r, size=22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(f"版本：V3.1    生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    set_run_font(r, size=11)

    doc.add_paragraph("")

    add_heading(doc, "1. 项目概述", level=1)
    add_para(doc, "Herbal Vision 是一个面向中草药识别与流媒体智能监控场景的 Web 控制台，支持图片/视频检测、模型切换、摄像机接入、目录监控、运行日志查看与草药知识联动展示。")
    add_para(doc, "系统采用前后端分离思路，前端负责统一操作入口与结果展示，后端负责模型推理、结果落盘、日志与配置管理。")

    add_heading(doc, "2. 技术栈", level=1)
    add_heading(doc, "2.1 前端", level=2)
    add_bullets(doc, [
        "HTML5 / CSS3 / 原生 JavaScript",
        "单页式多区域布局（仪表盘、检测中心、模型对比、监控中心、日志中心、参数设置）",
        "响应式卡片 UI、侧边栏导航、顶部操作区、结果区复用",
    ])

    add_heading(doc, "2.2 后端", level=2)
    add_bullets(doc, [
        "Python 3.10",
        "FastAPI + Uvicorn：提供 REST API 与 Web 页面入口",
        "Pydantic：请求数据建模与参数校验",
        "PyYAML：配置文件加载与持久化",
    ])

    add_heading(doc, "2.3 AI / 视觉与多媒体", level=2)
    add_bullets(doc, [
        "Ultralytics YOLO：目标检测与模型切换",
        "PyTorch：模型推理运行时",
        "OpenCV：图片/视频读取、结果绘制、视频编解码",
        "Pillow：中文检测框文本渲染",
        "httpx：调用外部 AI 接口生成草药知识与模型分析",
    ])

    add_heading(doc, "2.4 运行与部署", level=2)
    add_bullets(doc, [
        "Docker / docker-compose：容器化部署",
        "watchdog：目录监控自动识别",
        "onvif-zeep：ONVIF 摄像机解析",
    ])

    add_heading(doc, "3. 系统架构", level=1)
    add_para(doc, "系统结构可分为四层：前端展示层、API 接入层、业务服务层和数据/文件层。")
    add_bullets(doc, [
        "前端：负责页面渲染、交互、结果展示和参数配置",
        "API 层：提供认证、检测、模型管理、对比、监控、日志、设置等接口",
        "服务层：封装 YOLO 推理、草药知识联动、日志处理、摄像机接入与目录监控",
        "数据层：使用本地目录保存上传文件、输出结果、日志与配置文件",
    ])

    add_heading(doc, "4. 核心功能模块", level=1)
    add_heading(doc, "4.1 仪表盘", level=2)
    add_para(doc, "用于展示系统状态、当前模型、模型数量、目录监控处理数，并提供刷新数据与快捷跳转。")

    add_heading(doc, "4.2 检测中心", level=2)
    add_para(doc, "支持图片或视频统一上传；单图/单视频直接展示结果，多图时复用同一个结果窗口，通过上一项/下一项滑动查看。")
    add_para(doc, "右侧展示 AI 联动的草药知识，支持多草药折叠显示。")

    add_heading(doc, "4.3 模型对比", level=2)
    add_para(doc, "同一批图片在两个模型下分别识别，对比结果采用左右分栏展示；历史记录仅展示一条卡片，并支持点击加载详情与左右切换。")

    add_heading(doc, "4.4 监控中心", level=2)
    add_para(doc, "支持 RTSP / ONVIF 摄像机接入、实时流识别预览与摄像机状态管理。")

    add_heading(doc, "4.5 日志中心", level=2)
    add_para(doc, "统一展示系统日志、运行时间、来源模块与具体操作内容，便于追踪任务执行过程。")

    add_heading(doc, "4.6 参数设置", level=2)
    add_para(doc, "支持置信度、IoU、推理尺寸等检测参数调整，并新增 AI 配置窗口，用于填写 API 地址、模型名称、密钥与超时时间。")

    add_heading(doc, "5. UI 设计思路", level=1)
    add_bullets(doc, [
        "采用左侧固定侧边栏 + 右侧内容区的桌面控制台布局，减少页面跳转成本。",
        "整体使用浅色玻璃拟态与卡片化视觉，增强层次感和专业感。",
        "顶部区域集中放置版本号、刷新、跳转、账户、主题切换等高频操作，提升可达性。",
        "结果区、历史区、日志区等内容采用统一卡片风格，保持视觉一致。",
        "检测中心采用单结果窗口复用与批量滑动，避免多卡片堆叠造成页面过长。",
        "模型对比页采用单条历史卡片 + 滑动切换 + 点击加载详情，兼顾紧凑性与可追溯性。",
        "参数设置页将 AI 配置与检测参数分区，避免将外部服务配置散落到 compose 中。",
    ])

    add_heading(doc, "6. 关键数据流", level=1)
    add_bullets(doc, [
        "检测上传：前端上传文件 → 后端保存 → YOLO 推理 → 绘制结果 → 返回结果 URL 与结构化 JSON。",
        "批量识别：前端上传多图 → 后端逐张识别 → 返回 items 数组 → 前端单卡滑动展示。",
        "模型对比：同批图片两模型推理 → 保存 meta.json 与结果图 → 历史记录页按 compare_id 加载详情。",
        "AI 联动：识别出草药名称后调用外部 AI 接口 → 返回草药介绍 / 对比建议 → 前端按单项或折叠形式展示。",
        "日志流：业务服务输出结构化日志 → 日志文件读取解析 → 前端日志中心按时间/来源/内容展示。",
    ])

    add_heading(doc, "7. 配置与部署设计", level=1)
    add_para(doc, "系统默认通过 config/config.yaml 维护基础配置，运行时支持环境变量覆盖；AI 配置已改为在 Web UI 的参数设置页中保存，避免在 docker-compose 中直接硬编码。")
    add_bullets(doc, [
        "Docker Compose 负责基础容器启动、卷挂载与资源限制。",
        "后台启动时自动创建 uploads、outputs、watch、logs 等目录。",
        "AI 配置通过 /api/settings/ai 写入 config/config.yaml，重启后保持持久化。",
    ])

    add_heading(doc, "8. 设计亮点", level=1)
    add_bullets(doc, [
        "将 AI 配置从 compose 中剥离，改为前端可视化配置，更适合交付与运维。",
        "通过统一的卡片化 UI 与收纳式结果区，控制信息密度。",
        "多图识别、历史记录与日志中心均采用“单条聚焦 + 详情加载”的思路。",
        "中文标签绘制、日志解析、历史追踪等环节均考虑了可读性与可维护性。",
    ])

    add_heading(doc, "9. 总结", level=1)
    add_para(doc, "Herbal Vision 已形成从上传识别、模型对比、摄像机监控、目录监控到日志追踪的完整工作流，并通过 AI 配置窗口实现了外部智能服务的可视化管理。整体设计兼顾部署简洁性、功能可扩展性与桌面控制台的使用体验。")

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT_FILE)
    print(f"saved: {OUT_FILE.resolve()}")


if __name__ == "__main__":
    main()
